#!/usr/bin/env python3
from __future__ import annotations

import argparse
import importlib.util
import tempfile
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("ERROR: python-docx is required") from exc

HERE = Path(__file__).resolve().parent
RENDERER = HERE / "render_report.py"
ROOT = HERE.parent
SPEC = ROOT / "assets/report-spec.example.json"
THEME = ROOT / "assets/report-theme.json"
TOKENS = {
    "{{DOCUMENT_TYPE}}",
    "{{DOCUMENT_ID}}",
    "{{DATE}}",
    "{{CONFIDENTIALITY}}",
    "{{REPORT_BODY}}",
}


def load_renderer():
    spec = importlib.util.spec_from_file_location("dk_sport_renderer", RENDERER)
    if spec is None or spec.loader is None:
        raise RuntimeError("cannot load sport DOCX renderer")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def iter_paragraphs(part):
    yield from part.paragraphs
    for table in part.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def all_text(doc: Document) -> str:
    parts = [doc]
    for section in doc.sections:
        parts.extend([section.header, section.footer])
    return "\n".join(p.text for part in parts for p in iter_paragraphs(part))


def build_valid_template(path: Path) -> None:
    doc = Document()
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{DOCUMENT_TYPE}} · {{DOCUMENT_ID}}"
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "{{CONFIDENTIALITY}}"
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Datum: {{DATE}}"
    doc.add_paragraph("{{REPORT_BODY}}")
    doc.save(path)


def build_invalid_nested_body_template(path: Path) -> None:
    doc = Document()
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{DOCUMENT_TYPE}} · {{DOCUMENT_ID}} · {{DATE}}"
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "{{CONFIDENTIALITY}}"
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{REPORT_BODY}}"
    doc.save(path)


def verify() -> None:
    renderer = load_renderer()
    with tempfile.TemporaryDirectory(prefix="dk_template_override_contract_") as td:
        root = Path(td)
        valid = root / "valid.docx"
        rendered = root / "rendered.docx"
        invalid = root / "invalid.docx"
        invalid_output = root / "invalid-output.docx"

        build_valid_template(valid)
        renderer.render(SPEC, rendered, valid, THEME)
        result = Document(rendered)
        text = all_text(result)
        leftovers = sorted(token for token in TOKENS if token in text)
        if leftovers:
            raise RuntimeError("valid override left unresolved token(s): " + ", ".join(leftovers))
        for expected in ("Trainingsplan", "DK-ST-2026-001", "19.08.2026", "Vertraulich", "12-Wochen Trainingsplan"):
            if expected not in text:
                raise RuntimeError(f"valid override lost expected rendered value: {expected}")

        build_invalid_nested_body_template(invalid)
        try:
            renderer.render(SPEC, invalid_output, invalid, THEME)
        except ValueError as exc:
            message = str(exc)
            if "{{REPORT_BODY}}" not in message or "top-level" not in message:
                raise RuntimeError(f"invalid override failed for unexpected reason: {message}") from exc
        else:
            raise RuntimeError("invalid override with nested {{REPORT_BODY}} was accepted")


def main() -> int:
    parser = argparse.ArgumentParser(description="Verify the controlled Dr. Komorowski DOCX template override contract.")
    parser.parse_args()
    try:
        verify()
    except (RuntimeError, ValueError) as exc:
        print(f"ERROR: {exc}")
        return 1
    print("template override contract: PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
