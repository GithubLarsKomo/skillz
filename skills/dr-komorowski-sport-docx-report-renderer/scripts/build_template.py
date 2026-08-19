#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import tempfile
from pathlib import Path

try:
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("ERROR: build requires python-docx and Pillow") from exc

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
ASSET = ROOT / "assets"
DEFAULT_OUTPUT = ASSET / "dr-komorowski-report-template.docx"
DEFAULT_SNAPSHOT = ASSET / "dr-komorowski-report-template.docx.b64"
NAVY = "173652"
DARK = "1C2B3A"
BODY = "24313E"
TEAL = "2B8884"
TEAL_TEXT = "246F6C"
MUTED = "6B7785"
BORDER = "D6E0E6"


def _font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def _logo_png(path: Path) -> None:
    image = Image.new("RGBA", (2600, 640), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((25, 80, 560, 600), radius=85, outline="#" + NAVY, width=25)
    draw.text((115, 185), "DK", font=_font(270, True), fill="#" + NAVY)
    for y, length in [(205, 235), (325, 330), (445, 185)]:
        draw.rounded_rectangle((610, y, 610 + length, y + 34), radius=14, fill="#" + TEAL)
    draw.text((1010, 200), "DR. KOMOROWSKI", font=_font(145, True), fill="#" + NAVY)
    draw.text((1010, 385), "SPORTDIAGNOSE UND TRAININGSZENTRUM", font=_font(74, True), fill="#" + TEAL_TEXT)
    image.save(path, dpi=(300, 300))


def _border(cell, edge: str, color: str, size: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    for key, value in {"val": "single", "sz": str(size), "color": color, "space": "0"}.items():
        node.set(qn("w:" + key), value)


def _style(doc: Document, name: str, size: float, color: str, bold: bool = False):
    styles = doc.styles
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH) if name not in styles else styles[name]
    style.font.name = "Aptos"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    return style


def build(output: Path = DEFAULT_OUTPUT) -> Path:
    ASSET.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="dk_sport_template_") as td:
        logo = Path(td) / "logo.png"
        _logo_png(logo)
        doc = Document()
        sec = doc.sections[0]
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.left_margin, sec.right_margin = Mm(18), Mm(18)
        sec.top_margin, sec.bottom_margin = Mm(22), Mm(18)
        sec.header_distance, sec.footer_distance = Mm(7), Mm(7)

        normal = doc.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(BODY)
        normal.paragraph_format.space_after = Pt(5)
        for name, size, color, bold, before, after in [
            ("Title", 25, DARK, False, 0, 10),
            ("Heading 1", 17, NAVY, False, 16, 7),
            ("Heading 2", 12.5, TEAL_TEXT, True, 12, 5),
            ("Heading 3", 10.5, DARK, True, 9, 3),
        ]:
            st = doc.styles[name]
            st.font.name = "Aptos"
            st.font.size = Pt(size)
            st.font.bold = bold
            st.font.color.rgb = RGBColor.from_string(color)
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
        _style(doc, "DK Eyebrow", 8.5, TEAL_TEXT, True)
        _style(doc, "DK Subtitle", 11.5, MUTED)
        _style(doc, "DK Caption", 7.5, MUTED)
        _style(doc, "DK Callout", 9.5, BODY)

        header = sec.header
        header.paragraphs[0].clear()
        header.paragraphs[0].paragraph_format.space_after = Pt(0)
        table = header.add_table(rows=1, cols=2, width=Mm(174))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left, right = table.rows[0].cells
        left.width, right.width = Mm(112), Mm(62)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _border(cell, "bottom", BORDER, 7)
        left.paragraphs[0].paragraph_format.space_after = Pt(3)
        left.paragraphs[0].add_run().add_picture(str(logo), width=Mm(72))
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        for idx, token in enumerate(["{{DOCUMENT_TYPE}}", "{{DOCUMENT_ID}}", "{{DATE}}"]):
            run = p.add_run(token)
            run.font.name = "Aptos"
            run.font.size = Pt(8.2 if idx == 0 else 7.5)
            run.font.bold = idx == 0
            run.font.color.rgb = RGBColor.from_string(NAVY if idx == 0 else MUTED)
            if idx < 2:
                run.add_break()

        footer = sec.footer
        footer.paragraphs[0].clear()
        ftable = footer.add_table(rows=1, cols=2, width=Mm(174))
        ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
        ftable.autofit = False
        lf, rf = ftable.rows[0].cells
        lf.width, rf.width = Mm(120), Mm(54)
        for cell in (lf, rf):
            _border(cell, "top", BORDER, 5)
        p = lf.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Dr. Komorowski Sportdiagnose und Trainingszentrum")
        r.font.name = "Aptos"
        r.font.size = Pt(6.8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        p = rf.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("{{CONFIDENTIALITY}}")
        r.font.name = "Aptos"
        r.font.size = Pt(6.8)
        r.font.color.rgb = RGBColor.from_string(MUTED)

        body = doc.add_paragraph("{{REPORT_BODY}}")
        body.style = "Normal"
        props = doc.core_properties
        props.title = "Dr. Komorowski Sport Report Template"
        props.subject = "Reusable DOCX template for sport diagnostics and training reports"
        props.author = "Dr. Komorowski Sportdiagnose und Trainingszentrum"
        props.keywords = "sport diagnostics, training, report, template"
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
    return output


def write_snapshot(snapshot: Path = DEFAULT_SNAPSHOT) -> Path:
    with tempfile.TemporaryDirectory(prefix="dk_sport_snapshot_") as td:
        docx = Path(td) / "template.docx"
        build(docx)
        snapshot.parent.mkdir(parents=True, exist_ok=True)
        snapshot.write_text(base64.b64encode(docx.read_bytes()).decode("ascii") + "\n", encoding="ascii")
    return snapshot


def main() -> int:
    ap = argparse.ArgumentParser(description="Build the Dr. Komorowski sport DOCX template or its text-safe Base64 snapshot.")
    ap.add_argument("--output", type=Path, default=None)
    ap.add_argument("--snapshot", type=Path, default=None)
    args = ap.parse_args()
    if args.snapshot:
        print(write_snapshot(args.snapshot))
    else:
        print(build(args.output or DEFAULT_OUTPUT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
