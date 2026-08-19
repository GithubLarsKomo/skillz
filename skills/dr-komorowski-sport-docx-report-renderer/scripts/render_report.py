#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import io
import json
import math
import tempfile
from pathlib import Path
from typing import Any, Iterable

try:
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("ERROR: python-docx and Pillow are required") from exc

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = ROOT / "assets/dr-komorowski-report-template.docx"
DEFAULT_TEMPLATE_B64 = ROOT / "assets/dr-komorowski-report-template.docx.b64"
DEFAULT_THEME = ROOT / "assets/report-theme.json"
TOKENS = {"{{DOCUMENT_TYPE}}", "{{DOCUMENT_ID}}", "{{DATE}}", "{{CONFIDENTIALITY}}", "{{REPORT_BODY}}"}


def load_json(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"cannot read JSON {path}: {exc}") from exc
    if not isinstance(value, dict):
        raise ValueError(f"top-level JSON must be an object: {path}")
    return value


def _finite(value: Any) -> float:
    if isinstance(value, bool):
        raise ValueError("boolean is not a numeric chart value")
    number = float(value)
    if not math.isfinite(number):
        raise ValueError("chart values must be finite")
    return number


def validate(spec: dict[str, Any]) -> None:
    meta = spec.get("metadata")
    if not isinstance(meta, dict):
        raise ValueError("metadata object is required")
    for key in ("title", "date", "document_type"):
        if not isinstance(meta.get(key), str) or not meta[key].strip():
            raise ValueError(f"metadata.{key} is required")
    sections = spec.get("sections")
    if not isinstance(sections, list):
        raise ValueError("sections must be an array")
    allowed = {"paragraph", "heading", "subheading", "bullets", "table", "callout", "chart", "image", "spacer", "pagebreak"}
    for si, section in enumerate(sections):
        if not isinstance(section, dict) or not isinstance(section.get("blocks", []), list):
            raise ValueError(f"sections[{si}] must contain blocks[]")
        for bi, block in enumerate(section.get("blocks", [])):
            path = f"sections[{si}].blocks[{bi}]"
            if not isinstance(block, dict) or block.get("type") not in allowed:
                raise ValueError(f"unsupported block at {path}")
            if block["type"] == "table":
                cols, rows = block.get("columns"), block.get("rows")
                if not isinstance(cols, list) or not cols or not isinstance(rows, list):
                    raise ValueError(f"{path} requires columns and rows")
                if any(not isinstance(row, list) or len(row) != len(cols) for row in rows):
                    raise ValueError(f"{path} row width mismatch")
            if block["type"] == "image" and not isinstance(block.get("path"), str):
                raise ValueError(f"{path}.path is required")
            if block["type"] == "chart":
                if block.get("chart_type") != "lactate_hr_power":
                    raise ValueError(f"{path}.chart_type unsupported")
                data = block.get("data")
                if not isinstance(data, list) or len(data) < 2:
                    raise ValueError(f"{path}.data requires at least two points")
                powers = []
                for point in data:
                    if not isinstance(point, dict):
                        raise ValueError(f"{path}.data point must be an object")
                    powers.append(_finite(point.get("power_w")))
                    _finite(point.get("lactate_mmol_l"))
                    _finite(point.get("hr_bpm"))
                if any(b <= a for a, b in zip(powers, powers[1:])):
                    raise ValueError(f"{path}.power_w must be strictly increasing")
                for band in block.get("threshold_bands", []):
                    start, end = _finite(band.get("from_w")), _finite(band.get("to_w"))
                    if end <= start:
                        raise ValueError(f"{path} threshold band requires to_w > from_w")
                    if band.get("working_w") is not None:
                        working = _finite(band.get("working_w"))
                        if not start <= working <= end:
                            raise ValueError(f"{path}.working_w must lie inside threshold band")


def paragraphs(part) -> Iterable:
    yield from part.paragraphs
    for table in part.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs(cell)


def replace_tokens(doc, meta: dict[str, Any]) -> None:
    counts = {token: 0 for token in TOKENS}
    parts = [doc]
    for section in doc.sections:
        parts.extend([section.header, section.footer])
    for part in parts:
        for p in paragraphs(part):
            for token in TOKENS:
                if token in p.text:
                    counts[token] += 1
    missing = [token for token, count in counts.items() if not count]
    if missing:
        raise ValueError("template missing required token(s): " + ", ".join(sorted(missing)))
    values = {
        "{{DOCUMENT_TYPE}}": meta.get("document_type", ""),
        "{{DOCUMENT_ID}}": meta.get("document_id", ""),
        "{{DATE}}": meta.get("date", ""),
        "{{CONFIDENTIALITY}}": meta.get("confidentiality", "Vertraulich"),
    }
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in paragraphs(part):
                for token, value in values.items():
                    if token in p.text:
                        text = p.text.replace(token, str(value))
                        if p.runs:
                            p.runs[0].text = text
                            for run in p.runs[1:]:
                                run.text = ""
                        else:
                            p.add_run(text)


def marker(doc):
    for p in doc.paragraphs:
        if "{{REPORT_BODY}}" in p.text:
            return p
    raise ValueError("template is missing {{REPORT_BODY}}")


def add_p(doc, mark, text: str = "", style: str = "Normal"):
    p = doc.add_paragraph(text)
    try:
        p.style = style
    except KeyError:
        p.style = "Normal"
    mark._p.addprevious(p._p)
    return p


def add_t(doc, mark, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    mark._p.addprevious(table._tbl)
    return table


def shade(cell, fill: str) -> None:
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill.replace("#", ""))


def border(cell, edge: str, color: str, size: int = 4) -> None:
    pr = cell._tc.get_or_add_tcPr()
    borders = pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None:
        node = OxmlElement("w:" + edge)
        borders.append(node)
    for key, value in {"val": "single", "sz": str(size), "color": color.replace("#", "")}.items():
        node.set(qn("w:" + key), value)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def metadata_table(doc, mark, meta: dict[str, Any], theme: dict[str, Any]) -> None:
    fields = [
        ("Athlet", "athlete"), ("Sport", "sport"), ("Test / Phase", "test_or_phase"),
        ("Autor", "author"), ("Dokument-ID", "document_id"), ("Version", "version"), ("Datum", "date"),
    ]
    rows = [(label, meta.get(key)) for label, key in fields if isinstance(meta.get(key), str) and meta[key].strip()]
    if not rows:
        return
    table = add_t(doc, mark, len(rows), 2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        a, b = table.rows[i].cells
        a.width, b.width = Mm(36), Mm(138)
        a.text, b.text = label, value
        a.paragraphs[0].runs[0].bold = True
        shade(a, theme["colors"]["table_fill"])
        keep_row_together(table.rows[i])
        for cell in (a, b):
            border(cell, "bottom", theme["colors"]["border"])
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    add_p(doc, mark).paragraph_format.space_after = Pt(3)


def callout(doc, mark, block: dict[str, Any], theme: dict[str, Any]) -> None:
    colors = theme["colors"]
    kind = block.get("kind", "info")
    mapping = {
        "info": (colors["callout_fill"], colors["teal"]),
        "decision": (colors["table_fill"], colors["navy"]),
        "warning": (colors["warning_fill"], colors["warning_border"]),
        "neutral": (colors["callout_fill"], colors["border"]),
    }
    fill, line = mapping.get(kind, mapping["info"])
    table = add_t(doc, mark, 1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell, "left", line, 18)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    try:
        p.style = "DK Callout"
    except KeyError:
        pass
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(4)
    if block.get("title"):
        run = p.add_run(str(block["title"]).strip() + "\n")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(line.replace("#", ""))
    p.add_run(str(block.get("text", "")))
    add_p(doc, mark).paragraph_format.space_after = Pt(2)


def _font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def render_chart(block: dict[str, Any], theme: dict[str, Any], output: Path) -> None:
    data = block["data"]
    powers = [_finite(p["power_w"]) for p in data]
    lactate = [_finite(p["lactate_mmol_l"]) for p in data]
    hr = [_finite(p["hr_bpm"]) for p in data]
    width, height = 1800, 1000
    left, right, top, bottom = 180, 180, 125, 150
    plot_l, plot_r, plot_t, plot_b = left, width - right, top, height - bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    colors = theme["colors"]
    navy, teal = colors["navy"], colors["teal"]
    border_c, muted = colors["border"], colors["muted"]
    pmin, pmax = min(powers), max(powers)
    span = max(1.0, pmax - pmin)
    pmin -= span * 0.05
    pmax += span * 0.05
    lmax = max(2.0, math.ceil(max(lactate) * 1.15 * 2) / 2)
    hrmin = max(40.0, math.floor((min(hr) - 10) / 10) * 10)
    hrmax = min(230.0, math.ceil((max(hr) + 10) / 10) * 10)
    if hrmax <= hrmin:
        hrmax = hrmin + 20

    def x(v): return plot_l + (v - pmin) / (pmax - pmin) * (plot_r - plot_l)
    def yl(v): return plot_b - v / lmax * (plot_b - plot_t)
    def yh(v): return plot_b - (v - hrmin) / (hrmax - hrmin) * (plot_b - plot_t)

    for band in block.get("threshold_bands", []):
        start, end = _finite(band["from_w"]), _finite(band["to_w"])
        kind = str(band.get("kind", "lt1")).lower()
        fill = "#E8F5F3" if kind == "lt1" else "#FFF1CC"
        draw.rectangle((x(start), plot_t, x(end), plot_b), fill=fill)
        label = str(band.get("label", kind.upper()))
        draw.text((x(start) + 8, plot_t + 8), label, font=_font(28, True), fill=teal if kind == "lt1" else colors["warning_border"])
        if band.get("working_w") is not None:
            wx = x(_finite(band["working_w"]))
            for yy in range(plot_t, plot_b, 24):
                draw.line((wx, yy, wx, min(yy + 12, plot_b)), fill=teal if kind == "lt1" else colors["warning_border"], width=3)

    for i in range(6):
        y = plot_t + i * (plot_b - plot_t) / 5
        draw.line((plot_l, y, plot_r, y), fill=border_c, width=2)
    draw.line((plot_l, plot_t, plot_l, plot_b), fill=colors["dark"], width=3)
    draw.line((plot_r, plot_t, plot_r, plot_b), fill=colors["dark"], width=3)
    draw.line((plot_l, plot_b, plot_r, plot_b), fill=colors["dark"], width=3)

    lpts = [(x(p), yl(v)) for p, v in zip(powers, lactate)]
    hpts = [(x(p), yh(v)) for p, v in zip(powers, hr)]
    draw.line(lpts, fill=navy, width=8)
    draw.line(hpts, fill=teal, width=8)
    for px, py in lpts:
        draw.ellipse((px - 10, py - 10, px + 10, py + 10), fill=navy)
    for px, py in hpts:
        draw.ellipse((px - 10, py - 10, px + 10, py + 10), fill=teal)

    title = str(block.get("title", "Laktat und Herzfrequenz über Leistung"))
    draw.text((plot_l, 35), title, font=_font(42, True), fill=colors["dark"])
    draw.text((plot_l, height - 70), "Leistung (W)", font=_font(28, True), fill=colors["dark"])
    draw.text((25, plot_t - 10), "Laktat\n(mmol/L)", font=_font(25, True), fill=navy)
    draw.text((plot_r + 30, plot_t - 10), "Herzfrequenz\n(bpm)", font=_font(25, True), fill=teal)
    for p in powers:
        px = x(p)
        draw.line((px, plot_b, px, plot_b + 10), fill=colors["dark"], width=2)
        draw.text((px - 25, plot_b + 20), f"{p:g}", font=_font(23), fill=muted)
    for i in range(6):
        lv = lmax * (5 - i) / 5
        hv = hrmin + (hrmax - hrmin) * (5 - i) / 5
        yy = plot_t + i * (plot_b - plot_t) / 5
        draw.text((70, yy - 14), f"{lv:.1f}", font=_font(22), fill=navy)
        draw.text((plot_r + 30, yy - 14), f"{hv:.0f}", font=_font(22), fill=teal)
    draw.line((plot_l + 750, 75, plot_l + 820, 75), fill=navy, width=8)
    draw.text((plot_l + 835, 56), "Laktat", font=_font(24, True), fill=navy)
    draw.line((plot_l + 990, 75, plot_l + 1060, 75), fill=teal, width=8)
    draw.text((plot_l + 1075, 56), "Herzfrequenz", font=_font(24, True), fill=teal)
    image.save(output, dpi=(300, 300))


def render_blocks(doc, mark, spec: dict[str, Any], theme: dict[str, Any], base: Path, temp: Path) -> None:
    meta = spec["metadata"]
    if meta.get("department"):
        add_p(doc, mark, meta["department"], "DK Eyebrow").paragraph_format.space_after = Pt(4)
    add_p(doc, mark, meta["title"], "Title")
    if meta.get("subtitle"):
        add_p(doc, mark, meta["subtitle"], "DK Subtitle").paragraph_format.space_after = Pt(10)
    metadata_table(doc, mark, meta, theme)
    if isinstance(spec.get("summary"), str) and spec["summary"].strip():
        callout(doc, mark, {"kind": "info", "title": "Kurzfazit", "text": spec["summary"]}, theme)
    chart_counter = 0
    for section in spec["sections"]:
        if section.get("title"):
            add_p(doc, mark, str(section["title"]), "Heading 1")
        for block in section.get("blocks", []):
            kind = block["type"]
            if kind == "paragraph":
                add_p(doc, mark, str(block.get("text", "")))
            elif kind == "heading":
                level = min(max(int(block.get("level", 2)), 1), 3)
                add_p(doc, mark, str(block.get("text", "")), f"Heading {level}")
            elif kind == "subheading":
                add_p(doc, mark, str(block.get("text", "")), "Heading 2")
            elif kind == "bullets":
                for item in block.get("items", []):
                    add_p(doc, mark, str(item), "List Bullet")
            elif kind == "callout":
                callout(doc, mark, block, theme)
            elif kind == "table":
                cols, rows = block["columns"], block["rows"]
                table = add_t(doc, mark, len(rows) + 1, len(cols))
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.autofit = False
                weights = block.get("widths") or [1] * len(cols)
                total = sum(float(x) for x in weights)
                widths = [174 * float(x) / total for x in weights]
                repeat_header(table.rows[0])
                keep_row_together(table.rows[0])
                for j, col in enumerate(cols):
                    c = table.rows[0].cells[j]
                    c.width = Mm(widths[j])
                    c.text = str(col)
                    shade(c, theme["colors"]["table_fill"])
                    if c.paragraphs[0].runs:
                        c.paragraphs[0].runs[0].bold = True
                for i, row in enumerate(rows, 1):
                    keep_row_together(table.rows[i])
                    for j, value in enumerate(row):
                        c = table.rows[i].cells[j]
                        c.width = Mm(widths[j])
                        c.text = str(value)
                        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for row in table.rows:
                    for c in row.cells:
                        border(c, "bottom", theme["colors"]["border"])
                        for p in c.paragraphs:
                            p.paragraph_format.space_after = Pt(1)
                add_p(doc, mark).paragraph_format.space_after = Pt(2)
            elif kind == "chart":
                chart_counter += 1
                chart_path = temp / f"chart-{chart_counter}.png"
                render_chart(block, theme, chart_path)
                width = float(block.get("width_mm", 166))
                if not math.isfinite(width) or not 80 <= width <= 174:
                    raise ValueError("chart width_mm must be 80..174")
                p = add_p(doc, mark)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(chart_path), width=Mm(width))
                if block.get("caption"):
                    add_p(doc, mark, str(block["caption"]), "DK Caption").alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "image":
                source = Path(block["path"])
                source = source if source.is_absolute() else (base / source).resolve()
                if not source.is_file():
                    raise ValueError(f"image not found: {source}")
                width = float(block.get("width_mm", 150))
                if not math.isfinite(width) or not 20 <= width <= 174:
                    raise ValueError("image width_mm must be 20..174")
                p = add_p(doc, mark)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(source), width=Mm(width))
                if block.get("caption"):
                    add_p(doc, mark, str(block["caption"]), "DK Caption").alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "spacer":
                add_p(doc, mark).paragraph_format.space_after = Pt(float(block.get("height_pt", 8)))
            elif kind == "pagebreak":
                add_p(doc, mark).add_run().add_break(WD_BREAK.PAGE)


def load_template(template: Path, temp: Path):
    if template.is_file():
        return Document(str(template))
    if template == DEFAULT_TEMPLATE and DEFAULT_TEMPLATE_B64.is_file():
        try:
            raw = base64.b64decode(DEFAULT_TEMPLATE_B64.read_text(encoding="ascii"), validate=True)
        except (OSError, ValueError) as exc:
            raise ValueError(f"invalid bundled template representation: {DEFAULT_TEMPLATE_B64}") from exc
        return Document(io.BytesIO(raw))
    if template == DEFAULT_TEMPLATE:
        try:
            import importlib.util
            module_path = ROOT / "scripts/build_template.py"
            spec = importlib.util.spec_from_file_location("dk_template_builder", module_path)
            if spec is None or spec.loader is None:
                raise ValueError("cannot load template builder")
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)
            generated = temp / "template.docx"
            module.build(generated)
            return Document(str(generated))
        except Exception as exc:
            raise ValueError(f"bundled template unavailable and builder failed: {exc}") from exc
    raise ValueError(f"template not found: {template}")


def render(spec_path: Path, output: Path, template: Path, theme_path: Path) -> None:
    spec = load_json(spec_path)
    validate(spec)
    theme = load_json(theme_path)
    with tempfile.TemporaryDirectory(prefix="dk_sport_docx_") as td:
        temp = Path(td)
        doc = load_template(template, temp)
        replace_tokens(doc, spec["metadata"])
        mark = marker(doc)
        render_blocks(doc, mark, spec, theme, spec_path.parent, temp)
        mark._p.getparent().remove(mark._p)
        props = doc.core_properties
        props.title = spec["metadata"]["title"]
        props.author = spec["metadata"].get("author") or "Dr. Komorowski Sportdiagnose und Trainingszentrum"
        props.subject = spec["metadata"].get("document_type", "Sportreport")
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))


def main() -> int:
    ap = argparse.ArgumentParser(description="Render a structured Dr. Komorowski sport report spec to DOCX.")
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    ap.add_argument("--theme", type=Path, default=DEFAULT_THEME)
    args = ap.parse_args()
    try:
        render(args.input, args.output, args.template, args.theme)
    except ValueError as exc:
        print(f"ERROR: {exc}")
        return 2
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
