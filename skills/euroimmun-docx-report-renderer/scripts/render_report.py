#!/usr/bin/env python3
from __future__ import annotations

import argparse, json, math
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("ERROR: python-docx is required (pip install python-docx)") from exc

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = ROOT / "assets/euroimmun-report-template.docx"
DEFAULT_THEME = ROOT / "assets/report-theme.json"
TOKENS = {"{{DOCUMENT_TYPE}}", "{{DOCUMENT_ID}}", "{{DATE}}", "{{CONFIDENTIALITY}}", "{{REPORT_BODY}}"}


def load_json(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"cannot read JSON {path}: {exc}") from exc
    if not isinstance(value, dict):
        raise ValueError(f"top-level JSON must be an object: {path}")
    return value


def validate(spec: dict[str, Any]) -> None:
    meta = spec.get("metadata")
    if not isinstance(meta, dict):
        raise ValueError("metadata object is required")
    for key in ("title", "date", "document_type"):
        if not isinstance(meta.get(key), str) or not meta[key].strip():
            raise ValueError(f"metadata.{key} is required")
    sections = spec.get("sections")
    if not isinstance(sections, list):
        raise ValueError("sections must be an array")
    allowed = {"paragraph", "heading", "bullets", "table", "callout", "image", "spacer", "pagebreak"}
    for si, section in enumerate(sections):
        if not isinstance(section, dict) or not isinstance(section.get("blocks", []), list):
            raise ValueError(f"sections[{si}] must contain blocks[]")
        for bi, block in enumerate(section.get("blocks", [])):
            path = f"sections[{si}].blocks[{bi}]"
            if not isinstance(block, dict) or block.get("type") not in allowed:
                raise ValueError(f"unsupported block at {path}")
            if block["type"] == "table":
                cols, rows = block.get("columns"), block.get("rows")
                if not isinstance(cols, list) or not cols or not isinstance(rows, list):
                    raise ValueError(f"{path} requires columns and rows")
                if any(not isinstance(row, list) or len(row) != len(cols) for row in rows):
                    raise ValueError(f"{path} row width mismatch")
            if block["type"] == "image" and not isinstance(block.get("path"), str):
                raise ValueError(f"{path}.path is required")


def paragraphs(part) -> Iterable:
    yield from part.paragraphs
    for table in part.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs(cell)


def replace_tokens(doc, meta: dict[str, Any]) -> None:
    counts = {token: 0 for token in TOKENS}
    parts = [doc]
    for section in doc.sections:
        parts += [section.header, section.footer]
    for part in parts:
        for p in paragraphs(part):
            for token in TOKENS:
                if token in p.text:
                    counts[token] += 1
    missing = [token for token, count in counts.items() if not count]
    if missing:
        raise ValueError("template missing required token(s): " + ", ".join(sorted(missing)))
    values = {
        "{{DOCUMENT_TYPE}}": meta.get("document_type", ""), "{{DOCUMENT_ID}}": meta.get("document_id", ""),
        "{{DATE}}": meta.get("date", ""), "{{CONFIDENTIALITY}}": meta.get("confidentiality", "EUROIMMUN Confidential")
    }
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in paragraphs(part):
                for token, value in values.items():
                    if token in p.text:
                        text = p.text.replace(token, str(value))
                        if p.runs:
                            p.runs[0].text = text
                            for run in p.runs[1:]: run.text = ""
                        else: p.add_run(text)


def marker(doc):
    for p in doc.paragraphs:
        if "{{REPORT_BODY}}" in p.text: return p
    raise ValueError("template is missing {{REPORT_BODY}}")


def add_p(doc, mark, text="", style="Normal"):
    p = doc.add_paragraph(text)
    try: p.style = style
    except KeyError: p.style = "Normal"
    mark._p.addprevious(p._p)
    return p


def add_t(doc, mark, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    mark._p.addprevious(table._tbl)
    return table


def shade(cell, fill: str):
    pr = cell._tc.get_or_add_tcPr(); node = pr.find(qn("w:shd"))
    if node is None: node = OxmlElement("w:shd"); pr.append(node)
    node.set(qn("w:fill"), fill.replace("#", ""))


def border(cell, edge: str, color: str, size: int):
    pr = cell._tc.get_or_add_tcPr(); borders = pr.first_child_found_in("w:tcBorders")
    if borders is None: borders = OxmlElement("w:tcBorders"); pr.append(borders)
    node = borders.find(qn("w:" + edge))
    if node is None: node = OxmlElement("w:" + edge); borders.append(node)
    for key, value in {"val":"single", "sz":str(size), "color":color.replace("#", "")}.items(): node.set(qn("w:" + key), value)


def metadata_table(doc, mark, meta, theme):
    rows = [(k, meta.get(v)) for k, v in [("Author","author"),("Department","department"),("Document ID","document_id"),("Version","version"),("Date","date")]]
    rows = [(k,v) for k,v in rows if isinstance(v,str) and v.strip()]
    if not rows: return
    table = add_t(doc, mark, len(rows), 2); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    for i,(label,value) in enumerate(rows):
        a,b = table.rows[i].cells; a.width,b.width = Mm(34),Mm(132); a.text,b.text = label,value
        a.paragraphs[0].runs[0].bold = True; shade(a, theme["colors"]["table_label_fill"])
        for c in (a,b): border(c,"bottom",theme["colors"]["border"],4); c.paragraphs[0].paragraph_format.space_after=Pt(0)
    add_p(doc,mark).paragraph_format.space_after=Pt(4)


def callout(doc, mark, block, theme):
    kind=block.get("kind","info"); colors=theme["colors"]
    fill,line = {"info":(colors["callout_fill"],colors["green"]),"decision":(colors["callout_fill"],colors["green"]),"warning":(colors["warning_fill"],colors["warning_border"]),"neutral":(colors["neutral_fill"],colors["border"])}[kind]
    table=add_t(doc,mark,1,1); table.alignment=WD_TABLE_ALIGNMENT.LEFT
    trpr=table.rows[0]._tr.get_or_add_trPr(); trpr.append(OxmlElement("w:cantSplit"))
    cell=table.cell(0,0); shade(cell,fill); border(cell,"left",line,18); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=cell.paragraphs[0]
    try: p.style="EI Callout"
    except KeyError: pass
    p.paragraph_format.space_before=p.paragraph_format.space_after=Pt(4)
    if block.get("title"):
        run=p.add_run(str(block["title"]).strip()+"\n"); run.bold=True; run.font.color.rgb=RGBColor.from_string(line.replace("#",""))
    p.add_run(str(block.get("text","")))
    add_p(doc,mark).paragraph_format.space_after=Pt(2)


def render_blocks(doc, mark, spec, theme, base: Path):
    meta=spec["metadata"]
    if meta.get("department"): add_p(doc,mark,meta["department"],"EI Eyebrow").paragraph_format.space_after=Pt(5)
    add_p(doc,mark,meta["title"],"Title")
    if meta.get("subtitle"): add_p(doc,mark,meta["subtitle"],"EI Subtitle").paragraph_format.space_after=Pt(12)
    metadata_table(doc,mark,meta,theme)
    if isinstance(spec.get("summary"),str) and spec["summary"].strip(): callout(doc,mark,{"kind":"info","title":"Executive summary","text":spec["summary"]},theme)
    for section in spec["sections"]:
        if section.get("title"): add_p(doc,mark,section["title"],"Heading 1")
        for block in section.get("blocks",[]):
            kind=block["type"]
            if kind=="paragraph": add_p(doc,mark,block.get("text",""))
            elif kind=="heading": add_p(doc,mark,block.get("text",""),f"Heading {min(max(int(block.get('level',2)),1),3)}")
            elif kind=="bullets":
                for item in block.get("items",[]): add_p(doc,mark,item,"List Bullet")
            elif kind=="callout": callout(doc,mark,block,theme)
            elif kind=="table":
                cols,rows=block["columns"],block["rows"]; table=add_t(doc,mark,len(rows)+1,len(cols)); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
                weights=block.get("widths") or [1]*len(cols); total=sum(float(x) for x in weights); widths=[166*float(x)/total for x in weights]
                for j,col in enumerate(cols):
                    c=table.rows[0].cells[j]; c.width=Mm(widths[j]); c.text=str(col); shade(c,theme["colors"]["green_light"]); c.paragraphs[0].runs[0].bold=True
                for i,row in enumerate(rows,1):
                    for j,value in enumerate(row): table.rows[i].cells[j].width=Mm(widths[j]); table.rows[i].cells[j].text=str(value); table.rows[i].cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
                for row in table.rows:
                    for c in row.cells: border(c,"bottom",theme["colors"]["border"],4)
                add_p(doc,mark).paragraph_format.space_after=Pt(2)
            elif kind=="image":
                source=Path(block["path"]); source=source if source.is_absolute() else (base/source).resolve()
                if not source.is_file(): raise ValueError(f"image not found: {source}")
                width=float(block.get("width_mm",150));
                if not math.isfinite(width) or not 20<=width<=170: raise ValueError("image width_mm must be 20..170")
                p=add_p(doc,mark); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(source),width=Mm(width))
                if block.get("caption"): add_p(doc,mark,block["caption"],"EI Caption").alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif kind=="spacer": add_p(doc,mark).paragraph_format.space_after=Pt(float(block.get("height_pt",8)))
            elif kind=="pagebreak": add_p(doc,mark).add_run().add_break(WD_BREAK.PAGE)


def render(spec_path: Path, output: Path, template: Path, theme_path: Path):
    spec=load_json(spec_path); validate(spec); theme=load_json(theme_path)
    if not template.is_file(): raise ValueError(f"template not found: {template}")
    doc=Document(str(template)); replace_tokens(doc,spec["metadata"]); mark=marker(doc); render_blocks(doc,mark,spec,theme,spec_path.parent)
    mark._p.getparent().remove(mark._p); props=doc.core_properties; props.title=spec["metadata"]["title"]; props.author=spec["metadata"].get("author") or "EUROIMMUN Medizinische Labordiagnostika AG"
    output.parent.mkdir(parents=True,exist_ok=True); doc.save(str(output))


def main() -> int:
    ap=argparse.ArgumentParser(description="Render a structured EUROIMMUN report spec to DOCX.")
    ap.add_argument("input",type=Path); ap.add_argument("output",type=Path); ap.add_argument("--template",type=Path,default=DEFAULT_TEMPLATE); ap.add_argument("--theme",type=Path,default=DEFAULT_THEME)
    args=ap.parse_args()
    try: render(args.input,args.output,args.template,args.theme)
    except ValueError as exc: print(f"ERROR: {exc}"); return 2
    print(args.output); return 0


if __name__ == "__main__": raise SystemExit(main())
