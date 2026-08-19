#!/usr/bin/env python3
from __future__ import annotations

import tempfile
from pathlib import Path

try:
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("ERROR: build requires python-docx and Pillow") from exc

HERE = Path(__file__).resolve().parent
SKILL_ROOT = HERE.parent
ASSET = SKILL_ROOT / "assets"
OUTPUT = ASSET / "euroimmun-report-template.docx"
GREEN = "#79C143"
BLACK = "#171717"
GREY = "#63666A"


def _font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def _make_logo_png(path: Path) -> None:
    width, height = 2400, 420
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    word_font, tag_font = _font(190), _font(84)
    x, y = 30, 55
    draw.text((x, y), "euroimmun", font=word_font, fill=BLACK)
    word_right = draw.textbbox((x, y), "euroimmun", font=word_font)[2]
    draw.text((x + 5, 275), "From Revvity", font=tag_font, fill=BLACK)
    mark_x = word_right + 80
    line_y1, line_y2, line_w, line_h = 145, 278, 430, 22
    draw.rounded_rectangle((mark_x, line_y1, mark_x + line_w, line_y1 + line_h), radius=8, fill=GREEN)
    draw.rounded_rectangle((mark_x, line_y2, mark_x + line_w, line_y2 + line_h), radius=8, fill=GREEN)
    radius = 44
    center_y = (line_y1 + line_h + line_y2) // 2
    for center_x in [mark_x + 65 + i * 95 for i in range(4)]:
        draw.ellipse((center_x - radius, center_y - radius, center_x + radius, center_y + radius), fill=GREEN)
    image.save(path, dpi=(300, 300))


def _set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, values in kwargs.items():
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in values.items():
            element.set(qn("w:" + key), str(value))


def build(output: Path = OUTPUT) -> Path:
    ASSET.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="euroimmun_template_") as td:
        logo_png = Path(td) / "logo.png"
        _make_logo_png(logo_png)
        doc = Document()
        sec = doc.sections[0]
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.top_margin, sec.bottom_margin = Mm(24), Mm(18)
        sec.left_margin, sec.right_margin = Mm(20), Mm(20)
        sec.header_distance, sec.footer_distance = Mm(8), Mm(8)

        styles = doc.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10)
        styles["Normal"].font.color.rgb = RGBColor.from_string("242424")
        styles["Normal"].paragraph_format.space_after = Pt(6)
        for style_name, size, color, bold, before, after in [
            ("Title", 26, BLACK, False, 0, 12),
            ("Heading 1", 18, BLACK, False, 18, 8),
            ("Heading 2", 13, GREEN, True, 14, 6),
            ("Heading 3", 11, BLACK, True, 10, 4),
        ]:
            style = styles[style_name]
            style.font.name = "Aptos"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
        title_ppr = styles["Title"].element.get_or_add_pPr()
        title_border = title_ppr.find(qn("w:pBdr"))
        if title_border is not None:
            title_ppr.remove(title_border)

        for name, size, color, bold in [
            ("EI Eyebrow", 9, GREEN, True), ("EI Subtitle", 12, GREY, False),
            ("EI Caption", 8, GREY, False), ("EI Callout", 10, BLACK, False),
            ("EI Table", 8, BLACK, False), ("EI Table Header", 8, BLACK, True),
        ]:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH) if name not in styles else styles[name]
            style.font.name = "Aptos"; style.font.size = Pt(size); style.font.bold = bold
            style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))

        header = sec.header; header.paragraphs[0].clear(); header.paragraphs[0].paragraph_format.space_after = Pt(0)
        table = header.add_table(rows=1, cols=2, width=Mm(170)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
        table.columns[0].width, table.columns[1].width = Mm(110), Mm(60); left, right = table.rows[0].cells
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_border(cell, bottom={"val":"single","sz":"10","color":GREEN.replace("#", ""),"space":"0"})
        left.paragraphs[0].paragraph_format.space_after = Pt(4); left.paragraphs[0].add_run().add_picture(str(logo_png), width=Mm(69))
        right_p = right.paragraphs[0]; right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; right_p.paragraph_format.space_after = Pt(0)
        for index, text in enumerate(["{{DOCUMENT_TYPE}}", "{{DOCUMENT_ID}}", "{{DATE}}"]):
            run = right_p.add_run(text); run.font.name = "Aptos"; run.font.size = Pt(8 if index else 8.5); run.font.bold = index == 0; run.font.color.rgb = RGBColor.from_string("555555")
            if index < 2: run.add_break()

        footer = sec.footer; footer.paragraphs[0].clear(); footer.paragraphs[0].paragraph_format.space_before = Pt(4)
        ftable = footer.add_table(rows=1, cols=2, width=Mm(170)); ftable.alignment = WD_TABLE_ALIGNMENT.CENTER; ftable.autofit = False
        ftable.columns[0].width, ftable.columns[1].width = Mm(118), Mm(52); left_f, right_f = ftable.rows[0].cells
        for cell in (left_f, right_f): _set_cell_border(cell, top={"val":"single","sz":"6","color":"D7DED5","space":"0"})
        p = left_f.paragraphs[0]; p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(3), Pt(0)
        run = p.add_run("EUROIMMUN Medizinische Labordiagnostika AG · Seekamp 31 · 23560 Lübeck"); run.font.name, run.font.size = "Aptos", Pt(6.7); run.font.color.rgb = RGBColor.from_string("6A6A6A")
        p = right_f.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(3), Pt(0)
        run = p.add_run("{{CONFIDENTIALITY}}"); run.font.name, run.font.size = "Aptos", Pt(6.7); run.font.color.rgb = RGBColor.from_string("6A6A6A")

        body = doc.add_paragraph("{{REPORT_BODY}}"); body.style = "Normal"
        props = doc.core_properties; props.title = "EUROIMMUN Report Template"; props.subject = "Reusable public-reference template based on the 2025+ EUROIMMUN visual identity"; props.author = "EUROIMMUN Medizinische Labordiagnostika AG"; props.keywords = "EUROIMMUN, report, template, Revvity"
        output.parent.mkdir(parents=True, exist_ok=True); doc.save(output)
    return output


if __name__ == "__main__":
    print(build())
